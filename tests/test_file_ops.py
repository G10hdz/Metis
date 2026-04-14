"""Tests for file reader, editor, deleter, and bash agents."""

import json
import os
import tempfile
from pathlib import Path
from unittest.mock import patch, MagicMock

import pytest

from src.graph.nodes import (
    MetisState,
    ROUTE_FILE,
    ROUTE_FILE_EDIT,
    ROUTE_FILE_DELETE,
    ROUTE_BASH,
    ROUTE_GENERAL,
    router,
    _extract_path_from_query,
    _validate_file_path,
    _extract_file_content,
    _extract_pdf,
    _extract_docx,
    file_reader_agent,
    file_editor_agent,
    file_deleter_agent,
    bash_agent,
)
from src.config import settings


# ---- NLP Path extraction tests ----

class TestPathExtraction:
    @pytest.mark.parametrize("query,expected", [
        # Direct paths
        ("Lee el archivo ~/projects/Metis/src/config.py", "~/projects/Metis/src/config.py"),
        ("Muestra ~/projects/study-materials/data.json", "~/projects/study-materials/data.json"),
        ("Abre ./src/main.py", "./src/main.py"),
        ("Hola, ¿cómo estás?", None),
        ("Genera código en Python", None),
        # Trailing slash handling
        ("Lee ~/projects/Metis/src/", "~/projects/Metis/src"),
        ("Abre ~/projects/", "~/projects"),
        # Natural language with stopwords
        ("Lee el archivo config.py en la ruta ~/projects/Metis", "~/projects/Metis"),
        ("Muestrame el contenido del archivo ~/projects/test.txt", "~/projects/test.txt"),
        ("Por favor abre el documento settings.py en ~/projects/app", "~/projects/app"),
        ("Puedes mostrar el archivo readme.md en la carpeta ~/projects/docs", "~/projects/docs"),
    ])
    def test_extract_path(self, query, expected):
        result = _extract_path_from_query(query)
        assert result == expected, f"Query: {query!r}"

    def test_extract_path_with_spaces_and_quotes_unam_pdf(self, tmp_path):
        """Regression: extractor must keep spaced filenames and extensions."""
        target = tmp_path / "3. CUADERNO DE EJERCICIOS 2026.pdf"
        target.write_text("contenido", encoding="utf-8")
        expected = str(target)

        q_single = f"Lee el archivo '{target}'"
        q_double = f'Muestra el archivo "{target}".'
        q_unquoted = f"Abre {target} por favor"

        assert _extract_path_from_query(q_single) == expected
        assert _extract_path_from_query(q_double) == expected
        assert _extract_path_from_query(q_unquoted) == expected


# ---- Path validation tests ----

class TestPathValidation:
    def test_allowed_path(self, tmp_path):
        """A file under the allowed root should pass validation."""
        # Create a temp file under the allowed root
        allowed = Path.home() / "projects" / "_test_metis"
        allowed.mkdir(parents=True, exist_ok=True)
        test_file = allowed / "test.txt"
        test_file.write_text("hello", encoding="utf-8")

        resolved, error = _validate_file_path(str(test_file))
        assert error == ""
        assert resolved == test_file

        # Cleanup
        test_file.unlink()

    def test_path_outside_allowed(self):
        """A path outside the allowed root should be rejected."""
        resolved, error = _validate_file_path("/etc/passwd")
        assert "Access denied" in error
        assert resolved is None

    def test_nonexistent_file(self):
        """A file that doesn't exist should be rejected."""
        resolved, error = _validate_file_path("~/projects/nonexistent_file_xyz.txt")
        assert "not found" in error
        assert resolved is None

    def test_directory_not_file(self):
        """A directory should not be accepted as a file."""
        resolved, error = _validate_file_path("~/projects")
        assert "directory" in error
        assert resolved is None

    def test_binary_extension_rejected(self):
        """Binary file extensions should be rejected."""
        allowed = Path.home() / "projects" / "_test_metis"
        allowed.mkdir(parents=True, exist_ok=True)
        test_file = allowed / "image.png"
        test_file.write_bytes(b"\x89PNG\r\n\x1a\n")

        resolved, error = _validate_file_path(str(test_file))
        assert "not allowed" in error
        assert resolved is None

        test_file.unlink()


# ---- Router file/bash classification tests ----

class TestFileRouter:
    @pytest.mark.parametrize("query", [
        "Lee el archivo config.py en ~/projects/Metis",
        "Show me the file contents of ~/projects/Hypata/questions.json",
        "Muestra el archivo settings.py en ~/projects/test",
    ])
    def test_file_read_route(self, query):
        state = MetisState.from_query(query)
        result = router(state)
        assert result["route"] == ROUTE_FILE
        assert result["next_node"] == ROUTE_FILE

    @pytest.mark.parametrize("query", [
        "Edita config.py en ~/projects/Metis y cambia debug por True",
        "Modifica el archivo ~/projects/test.txt",
    ])
    def test_file_edit_route(self, query):
        state = MetisState.from_query(query)
        result = router(state)
        assert result["route"] == ROUTE_FILE_EDIT

    @pytest.mark.parametrize("query", [
        "Borra el archivo ~/projects/Metis/temp.txt",
        "Delete ~/projects/test/old_file.py",
    ])
    def test_file_delete_route(self, query):
        state = MetisState.from_query(query)
        result = router(state)
        assert result["route"] == ROUTE_FILE_DELETE

    @pytest.mark.parametrize("query", [
        "Ejecuta ls -la ~/projects/Metis",
        "Run the command grep -r test ~/Vscode-projects",
        "Shell command: find ~/Vscode-projects -name '*.py'",
    ])
    def test_bash_route(self, query):
        state = MetisState.from_query(query)
        result = router(state)
        assert result["route"] == ROUTE_BASH


# ---- File reader agent tests ----

class TestFileReaderAgent:
    def test_read_valid_file(self, tmp_path):
        """Reading a valid text file under the allowed root."""
        allowed = Path.home() / "projects" / "_test_metis"
        allowed.mkdir(parents=True, exist_ok=True)
        test_file = allowed / "readme.txt"
        test_file.write_text("Hello world\nLine 2\nLine 3", encoding="utf-8")

        state = MetisState(query=f"Lee {test_file}", file_path=str(test_file))
        result = file_reader_agent(state)

        assert result["next_node"] == "formatter"
        assert "Hello world" in result["response"]
        assert result["file_content"] == "Hello world\nLine 2\nLine 3"
        assert result["file_path"] == str(test_file)

        test_file.unlink()

    def test_read_nonexistent_file(self):
        """Reading a file that doesn't exist."""
        state = MetisState(query="Lee ~/projects/nonexistent.txt")
        result = file_reader_agent(state)

        assert "not found" in result["response"] or "No pude detectar" in result["response"]

    def test_read_no_path(self):
        """Query with no file path detected."""
        state = MetisState(query="Hola, ¿cómo estás?")
        result = file_reader_agent(state)

        assert "detectar" in result["response"].lower()


# ---- File editor agent tests ----

class TestFileEditorAgent:
    def test_edit_replace_flow(self, tmp_path):
        """Full replace flow: first call asks confirmation, second executes."""
        allowed = Path.home() / "projects" / "_test_metis"
        allowed.mkdir(parents=True, exist_ok=True)
        test_file = allowed / "edit_test.txt"
        test_file.write_text("foo bar baz", encoding="utf-8")

        # First call: asks for more info (no specific replace pattern)
        state = MetisState(query=f"Edita {test_file}", file_path=str(test_file))
        result = file_editor_agent(state)
        assert result["next_node"] == "formatter"
        assert "Reemplazar" in result["response"] or "cambio" in result["response"]

        # Cleanup
        test_file.write_text("foo bar baz", encoding="utf-8")  # restore
        test_file.unlink()


# ---- File deleter agent tests ----

class TestFileDeleterAgent:
    def test_delete_confirmation_flow(self, tmp_path):
        """Delete flow: first call asks confirmation."""
        allowed = Path.home() / "projects" / "_test_metis"
        allowed.mkdir(parents=True, exist_ok=True)
        test_file = allowed / "delete_me.txt"
        test_file.write_text("delete me", encoding="utf-8")

        # First call: asks for confirmation
        state = MetisState(query=f"Borra {test_file}", file_path=str(test_file))
        result = file_deleter_agent(state)

        assert result["awaiting_confirmation"] is True
        assert "eliminar" in result["response"].lower() or "seguro" in result["response"].lower()
        assert test_file.exists()  # file should still exist

        # Cleanup
        test_file.unlink()


# ---- Bash agent tests ----

class TestBashAgent:
    def test_allowed_command(self):
        """A whitelisted command should execute."""
        state = MetisState(query="ejecuta pwd")
        result = bash_agent(state)

        assert result["next_node"] == "formatter"
        assert "Command" in result["response"]
        assert result["bash_error"] == ""

    def test_dangerous_command(self):
        """A dangerous command should be blocked."""
        state = MetisState(query="run rm -rf /")
        result = bash_agent(state)

        assert "bloqueado" in result["response"] or "no permitido" in result["response"]

    def test_non_whitelisted_command(self):
        """A command not in the whitelist should be rejected."""
        state = MetisState(query="ejecuta python -c 'print(1)'")
        result = bash_agent(state)

        assert "no permitido" in result["response"] or "Comando" in result["response"]

    def test_empty_command(self):
        """Empty or command-only query should be rejected."""
        state = MetisState(query="ejecuta")
        result = bash_agent(state)

        assert "comando" in result["response"].lower() or "detecté" in result["response"]


# ---- Binary file extraction tests ----

class TestBinaryFileExtraction:
    def test_extract_plain_text(self, tmp_path):
        """Plain text files should be read directly."""
        test_file = tmp_path / "test.py"
        test_file.write_text("print('hello')", encoding="utf-8")
        content = _extract_file_content(test_file, ".py")
        assert content == "print('hello')"

    def test_extract_pdf_with_pymupdf(self, tmp_path):
        """PDF extraction via PyMuPDF (fitz)."""
        pytest.importorskip("fitz", reason="PyMuPDF not installed")

        # Create a simple PDF
        import fitz
        pdf_path = tmp_path / "test.pdf"
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((72, 72), "Hello from PDF\nLine 2")
        doc.save(str(pdf_path))
        doc.close()

        content = _extract_pdf(pdf_path)
        assert "Hello from PDF" in content
        assert "Line 2" in content

    def test_extract_docx(self, tmp_path):
        """DOCX extraction via python-docx."""
        pytest.importorskip("docx", reason="python-docx not installed")

        from docx import Document
        docx_path = tmp_path / "test.docx"
        doc = Document()
        doc.add_paragraph("Hello from DOCX")
        doc.add_paragraph("Second paragraph")
        doc.save(str(docx_path))

        content = _extract_docx(docx_path)
        assert "Hello from DOCX" in content
        assert "Second paragraph" in content

    def test_extract_docx_with_tables(self, tmp_path):
        """DOCX tables should be extracted as pipe-separated values."""
        pytest.importorskip("docx", reason="python-docx not installed")

        from docx import Document
        docx_path = tmp_path / "table_test.docx"
        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Name"
        table.cell(0, 1).text = "Value"
        table.cell(1, 0).text = "Foo"
        table.cell(1, 1).text = "42"
        doc.save(str(docx_path))

        content = _extract_docx(docx_path)
        assert "Name" in content
        assert "Foo" in content
        assert "42" in content

    def test_unsupported_binary_file(self, tmp_path):
        """Unsupported binary files should raise UnicodeDecodeError."""
        test_file = tmp_path / "image.png"
        test_file.write_bytes(b"\x89PNG\r\n\x1a\n\x00\x00")
        with pytest.raises(UnicodeDecodeError):
            _extract_file_content(test_file, ".png")

    def test_file_reader_pdf_integration(self, tmp_path):
        """file_reader_agent should handle PDF files end-to-end."""
        pytest.importorskip("fitz", reason="PyMuPDF not installed")

        # Create PDF under allowed root
        allowed = Path.home() / "projects" / "_test_metis"
        allowed.mkdir(parents=True, exist_ok=True)
        pdf_path = allowed / "test_read.pdf"

        import fitz
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((72, 72), "Metis PDF Test Content\nLine 2 of PDF")
        doc.save(str(pdf_path))
        doc.close()

        state = MetisState(query=f"Lee {pdf_path}", file_path=str(pdf_path))
        result = file_reader_agent(state)

        assert result["next_node"] == "formatter"
        assert "Metis PDF Test Content" in result["response"]
        assert result["file_path"] == str(pdf_path)

        pdf_path.unlink()

    def test_file_reader_docx_integration(self, tmp_path):
        """file_reader_agent should handle DOCX files end-to-end."""
        pytest.importorskip("docx", reason="python-docx not installed")

        # Create DOCX under allowed root
        allowed = Path.home() / "projects" / "_test_metis"
        allowed.mkdir(parents=True, exist_ok=True)
        docx_path = allowed / "test_read.docx"

        from docx import Document
        doc = Document()
        doc.add_paragraph("Metis DOCX Test Content")
        doc.add_paragraph("Second line of DOCX")
        doc.save(str(docx_path))

        state = MetisState(query=f"Lee {docx_path}", file_path=str(docx_path))
        result = file_reader_agent(state)

        assert result["next_node"] == "formatter"
        assert "Metis DOCX Test Content" in result["response"]
        assert result["file_path"] == str(docx_path)

        docx_path.unlink()

    def test_truncation_at_3500_chars(self, tmp_path):
        """Both PDF and DOCX content should be truncated at 3500 chars."""
        pytest.importorskip("docx", reason="python-docx not installed")

        # Create a large DOCX
        allowed = Path.home() / "projects" / "_test_metis"
        allowed.mkdir(parents=True, exist_ok=True)
        docx_path = allowed / "large.docx"

        from docx import Document
        doc = Document()
        # Add 5000 chars of content
        doc.add_paragraph("A" * 5000)
        doc.save(str(docx_path))

        state = MetisState(query=f"Lee {docx_path}", file_path=str(docx_path))
        result = file_reader_agent(state)

        # The response should contain truncation notice
        assert "truncated" in result["response"].lower()
        # file_content should have the full content
        assert len(result["file_content"]) > 3500

        docx_path.unlink()
