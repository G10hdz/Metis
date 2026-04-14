"""LangGraph node functions — router, rag_agent, code_agent, search_agent, general_agent, formatter.

CRITICAL: Every node returns a dict whose keys are a subset of MetisState fields.
LangGraph StateGraph merges these dicts into accumulated state. Never return raw strings.
"""

from __future__ import annotations

import ast
import json
import logging
import os
import re
import time
from pathlib import Path
from typing import Any

from src.config import settings
from src.config import ollama as ollama_cfg
from src.graph.state import MetisState
from src.memory.store import get_store
from src.telemetry.store import get_telemetry
from src.utils.vram_guard import vram_call_structured
from src.utils.fallback import call_with_fallback

logger = logging.getLogger(__name__)

# --- Route labels ---
ROUTE_RAG = "rag"
ROUTE_CODE = "code"
ROUTE_SEARCH = "search"
ROUTE_SEARCH_CONTINUE = "search_continue"
ROUTE_FILE = "file"
ROUTE_FILE_EDIT = "file_edit"
ROUTE_FILE_DELETE = "file_delete"
ROUTE_BASH = "bash"
ROUTE_GENERAL = "general"
ROUTE_PRACTICE = "pronunciation_practice"

# --- Router keyword maps ---
_CODE_KEYWORDS = {"code", "function", "script", "python", "debug", "error", "bug", "implement",
                  "refactor", "algorithm", "api", "endpoint", "class", "def", "import", "pip",
                  "install", "docker", "deploy", "server", "http", "json", "yaml",
                  "lambda", "terraform", "aws", "s3", "ec2"}

_RAG_KEYWORDS = {"what is", "how to", "explain", "describe", "define", "concept", "documentation",
                  "docs", "manual", "guide", "tutorial", "learn", "understand", "why", "when",
                  "where", "which", "history", "difference between"}

_SEARCH_KEYWORDS = {"search", "find", "latest", "news", "current", "recent", "what's new",
                    "trending", "2025", "2026", "update", "breaking", "announced", "released"}

_DEEPER_KEYWORDS = {"go deeper", "deeper", "more", "dig deeper", "elaborate", "tell me more",
                    "continue", "keep going", "more details", "dive deeper"}

_FILE_KEYWORDS = {"read", "open", "show", "display", "cat ", "show me",
                   "file contents", "content of", "contents of", "lee", "abre",
                   "muestra", "contenido de", "archivo", "mira", "ver"}

_FILE_EDIT_KEYWORDS = {"edit", "modify", "change", "replace", "update file", "fix in",
                        "edita", "modifica", "cambia", "reemplaza", "corrige", "agrega"}

_FILE_DELETE_KEYWORDS = {"delete", "remove", "erase", "trash", "borra", "elimina", "quita",
                          "borrar", "eliminar"}

_BASH_KEYWORDS = {"run", "execute", "bash", "shell", "command", "terminal",
                  "ejecuta", "corre", "comando", "terminal", "shell", "run the command"}

_PRACTICE_KEYWORDS = {"practice", "pronunciation", "pronounce", "read aloud", "speak",
                      "repeat after", "say this", "read this", "practicar", "pronunciación",
                      "pronunciar", "lee esto", "di esto", "repetir"}

# Regex patterns that look like a file path in a query
_PATH_PATTERNS = [
    r'(/home/\S+)',           # absolute Linux paths
    r'(~/\S+)',               # tilde paths
    r'(\./\S+)',              # relative paths
    r'(\w+[/\\]\w+[/\\]\S+)', # Windows-style or nested paths
]

# Stopwords to strip from natural language queries before path extraction
_FILE_STOPWORDS = {
    "lee", "abre", "muestra", "muestrame", "mostrar", "ver", "mira", "mirame",
    "el", "la", "los", "las", "un", "una", "unos", "unas",
    "archivo", "archivos", "file", "files", "documento", "doc",
    "de", "del", "al", "en", "con", "sin", "para", "por",
    "ruta", "path", "camino", "carpeta", "folder", "directorio",
    "contenido", "contenidos", "content", "contents",
    "el archivo", "la ruta", "este", "esta", "estos", "estas",
    "read", "open", "show", "display", "cat", "show me",
    "file contents", "content of", "contents of", "contenido de", "archivo",
    "please", "por favor", "puedes", "podes",
}


def _extract_path_from_query(query: str) -> str | None:
    """
    Extract a file path from natural language queries.

    Rules:
    1) If a path is quoted ('...' or "..."), extract full quoted content.
    2) If unquoted and path contains spaces, greedily rebuild candidates and
       prefer the longest one that exists in the filesystem.
    3) Preserve extensions like .pdf/.docx (avoid truncating before extension).
    """

    def _clean_candidate(raw: str) -> str:
        candidate = raw.strip()
        candidate = candidate.lstrip("([{\"'")
        candidate = candidate.rstrip(",;:!?)]}\"'")
        # Keep extension when user ends sentence with trailing dot: file.pdf.
        if candidate.endswith(".") and re.search(r"\.[A-Za-z0-9]{1,8}\.$", candidate):
            candidate = candidate[:-1]
        if candidate != "/" and len(candidate) > 1:
            candidate = candidate.rstrip("/")
        return candidate.strip()

    def _path_exists(raw: str) -> bool:
        try:
            return Path(raw).expanduser().exists()
        except OSError:
            return False

    def _looks_like_path(raw: str) -> bool:
        return (
            "/" in raw
            or "\\" in raw
            or raw.startswith(("~/", "./", "../"))
            or bool(re.search(r"\.[A-Za-z0-9]{1,8}$", raw))
        )

    def _greedy_from_tokens(tokens: list[str]) -> str | None:
        if not tokens:
            return None

        # Longest-prefix-first: captures unquoted paths with spaces.
        for end in range(len(tokens), 0, -1):
            candidate = _clean_candidate(" ".join(tokens[:end]))
            if candidate and _path_exists(candidate):
                return candidate

        # If no existing path matched, keep up to extension when present.
        joined = " ".join(tokens)
        ext_match = re.match(r"^(.+\.[A-Za-z0-9]{1,8})(?=$|\s|[,;:!?)]})", joined)
        if ext_match:
            candidate = _clean_candidate(ext_match.group(1))
            if candidate:
                return candidate

        first = _clean_candidate(tokens[0])
        if first and _looks_like_path(first):
            return first
        return None

    # 1) Quoted path extraction: keep everything between matching quotes.
    quoted_candidates = []
    for match in re.finditer(r"(['\"])(.+?)\1", query):
        quoted = _clean_candidate(match.group(2))
        if quoted:
            quoted_candidates.append(quoted)

    if quoted_candidates:
        for candidate in quoted_candidates:
            if _looks_like_path(candidate):
                return candidate
        for candidate in quoted_candidates:
            if _path_exists(candidate):
                return candidate
        return quoted_candidates[0]

    # 2) Unquoted greedy extraction from first explicit path starter.
    start_match = re.search(r"(?:(?<=\s)|^)(~/|\./|\.\./|/|[A-Za-z]:[\\/])", query)
    if start_match:
        start = start_match.start(1)
        greedy = _greedy_from_tokens(query[start:].split())
        if greedy:
            return greedy

    # 3) Token-level fallback for path-like fragments.
    tokens = query.split()
    for idx, token in enumerate(tokens):
        if "/" in token or "\\" in token or token.startswith(("~/", "./", "../")):
            greedy = _greedy_from_tokens(tokens[idx:])
            if greedy:
                return greedy

    # 4) Last-resort compact regex extraction (without spaces).
    for pattern in _PATH_PATTERNS:
        match = re.search(pattern, query)
        if match:
            candidate = _clean_candidate(match.group(1))
            if candidate:
                return candidate

    return None


def _telemetry_hook(state: MetisState, response: str, model: str, start: float, error: str = "") -> None:
    """Log conversation to SQLite telemetry store."""
    try:
        latency_ms = (time.monotonic() - start) * 1000
        get_telemetry().log(
            query=state.query,
            route=state.route,
            response=response[:2000],  # truncate for storage
            model=model,
            latency_ms=round(latency_ms, 1),
            error=error,
            source=state.source or "cli",
        )
    except Exception as exc:
        logger.debug("telemetry_hook failed: %s", exc)


def _validate_file_path(raw_path: str) -> tuple[Path | None, str]:
    """
    Validate a file path against security rules.
    Returns (resolved_path, error_message).
    If valid, error_message is empty.
    """
    # Resolve to absolute path
    if raw_path.startswith('~'):
        resolved = Path(raw_path).expanduser()
    else:
        resolved = Path(raw_path) if Path(raw_path).is_absolute() else Path.cwd() / Path(raw_path)

    resolved = resolved.resolve()

    # Check it's under the allowed root
    allowed_root = settings.ALLOWED_FILE_ROOT.resolve()
    try:
        resolved.relative_to(allowed_root)
    except ValueError:
        return None, f"⛔ Access denied. I can only read files under `{allowed_root}`."

    # Reject symlinks
    if resolved.is_symlink():
        return None, f"⛔ Symlinks are not allowed for security reasons."

    # Must exist
    if not resolved.exists():
        return None, f"📭 File not found: `{resolved}`"

    # Must be a file
    if not resolved.is_file():
        return None, f"📁 That's a directory, not a file: `{resolved.name}`"

    # Check extension (allow files without extension too, like Dockerfile, Makefile)
    ext = resolved.suffix.lower()
    name = resolved.name.lower()
    allowed_no_ext = {"dockerfile", "makefile", "license", "readme", "vagrantfile"}
    if ext and ext not in settings.ALLOWED_FILE_EXTENSIONS and name not in allowed_no_ext:
        return None, f"⛔ File type `{ext or '(no extension)'}` is not allowed. I can only read text files."

    # Check file size
    size = resolved.stat().st_size
    if size > settings.MAX_FILE_SIZE:
        return None, f"📏 File too large ({size / 1024:.1f}KB). I can only read files up to {settings.MAX_FILE_SIZE / 1024:.0f}KB."

    return resolved, ""


def router(state: MetisState) -> dict:
    """
    Classify query into rag | code | search | search_continue | file | file_edit | file_delete | bash | general.
    Returns dict compatible with StateGraph merge.
    """
    query_lower = state.query.lower().strip()
    if not query_lower:
        return {"next_node": "formatter", "route": ROUTE_GENERAL, "response": "Empty query."}

    # Check "go deeper" FIRST — user wants to continue a previous search
    deeper_score = sum(1 for kw in _DEEPER_KEYWORDS if kw in query_lower)
    if deeper_score > 0 and state.search_context:
        logger.info("router query='%s' route=%s (deeper)", state.query[:80], ROUTE_SEARCH_CONTINUE)
        return {"route": ROUTE_SEARCH_CONTINUE, "next_node": ROUTE_SEARCH_CONTINUE}

    # Check for research/vault building
    research_score = sum(1 for kw in _RESEARCH_KEYWORDS if kw in query_lower)
    if research_score >= 2 or (research_score >= 1 and any(kw in query_lower for kw in ("vault", "obsidian", "wiki", "nota", "note", "knowledge"))):
        logger.info("router query='%s' route=%s", state.query[:80], ROUTE_RESEARCH)
        return {"route": ROUTE_RESEARCH, "next_node": ROUTE_RESEARCH}

    # Check for confirmation response (delete/edit flow)
    if state.awaiting_confirmation and query_lower in ("yes", "si", "sí", "y", "confirm", "dale", "ok", "go ahead", "adelante"):
        logger.info("router query='%s' route=%s (confirmation)", state.query[:80], state.route)
        return {"route": state.route, "next_node": state.route, "awaiting_confirmation": False}
    if state.awaiting_confirmation and query_lower in ("no", "cancel", "stop", "nel", "chale"):
        logger.info("router query='%s' route=general (cancelled)", state.query[:80])
        return {"route": ROUTE_GENERAL, "next_node": ROUTE_GENERAL, "awaiting_confirmation": False,
                "response": "✅ Cancelado. No se hizo ningún cambio."}

    # Extract path if present
    extracted_path = _extract_path_from_query(state.query)

    # Score against keyword sets
    file_score = sum(1 for kw in _FILE_KEYWORDS if kw in query_lower)
    edit_score = sum(1 for kw in _FILE_EDIT_KEYWORDS if kw in query_lower)
    delete_score = sum(1 for kw in _FILE_DELETE_KEYWORDS if kw in query_lower)
    bash_score = sum(1 for kw in _BASH_KEYWORDS if kw in query_lower)
    code_score = sum(1 for kw in _CODE_KEYWORDS if kw in query_lower)
    rag_score = sum(1 for kw in _RAG_KEYWORDS if kw in query_lower)
    search_score = sum(1 for kw in _SEARCH_KEYWORDS if kw in query_lower)
    practice_score = sum(1 for kw in _PRACTICE_KEYWORDS if kw in query_lower)

    # Priority: delete > edit > bash > file (read) > practice > search > code > rag
    if extracted_path and delete_score > 0:
        logger.info("router query='%s' route=%s path='%s'", state.query[:80], ROUTE_FILE_DELETE, extracted_path)
        return {"route": ROUTE_FILE_DELETE, "next_node": ROUTE_FILE_DELETE, "file_path": extracted_path}

    if extracted_path and edit_score > 0:
        logger.info("router query='%s' route=%s path='%s'", state.query[:80], ROUTE_FILE_EDIT, extracted_path)
        return {"route": ROUTE_FILE_EDIT, "next_node": ROUTE_FILE_EDIT, "file_path": extracted_path}

    if extracted_path and file_score > 0:
        logger.info("router query='%s' route=%s path='%s'", state.query[:80], ROUTE_FILE, extracted_path)
        return {"route": ROUTE_FILE, "next_node": ROUTE_FILE, "file_path": extracted_path}

    if bash_score > 0 or (extracted_path and any(kw in query_lower for kw in ("run", "execute", "ejecuta", "corre"))):
        # If it looks like "run command X" or "execute ls", route to bash
        logger.info("router query='%s' route=%s", state.query[:80], ROUTE_BASH)
        return {"route": ROUTE_BASH, "next_node": ROUTE_BASH}

    if practice_score > 0:
        logger.info("router query='%s' route=%s", state.query[:80], ROUTE_PRACTICE)
        return {"route": ROUTE_PRACTICE, "next_node": ROUTE_PRACTICE}

    if search_score > 0 and search_score >= code_score and search_score >= rag_score:
        route = ROUTE_SEARCH
    elif code_score > 0 and code_score >= rag_score:
        route = ROUTE_CODE
    elif rag_score > 0:
        route = ROUTE_RAG
    else:
        route = ROUTE_GENERAL

    logger.info("router query='%s' route=%s code=%d rag=%d search=%d file=%d edit=%d delete=%d bash=%d practice=%d",
                state.query[:80], route, code_score, rag_score, search_score,
                file_score, edit_score, delete_score, bash_score, practice_score)

    return {
        "route": route,
        "next_node": route,
    }


def rag_agent(state: MetisState) -> dict:
    """
    Retrieve from ChromaDB and answer using context.
    If retrieval score < threshold, fall back to general agent.
    """
    start = time.monotonic()
    store = get_store()
    results = store.query(state.query, n_results=3)

    if not results or not results.get("documents"):
        logger.info("rag_agent empty retrieval, falling back to general")
        return _general_answer(state, start)

    docs = results["documents"]
    distances = results.get("distances", [[]])[0]

    context_parts: list[str] = []
    for i, (doc, dist) in enumerate(zip(docs, distances)):
        score = 1.0 - dist
        if score < settings.RAG_SCORE_THRESHOLD:
            continue
        context_parts.append(doc)

    if not context_parts:
        logger.info("rag_agent no docs above threshold, falling back to general")
        return _general_answer(state, start)

    kb_context = "\n\n---\n\n".join(context_parts)
    system_prompt = (
        "You are Metis, a helpful AI assistant. Answer based on the provided context. "
        "If the context doesn't fully answer the question, say so and provide your best answer."
    )
    prompt = f"Context:\n{kb_context}\n\nQuestion: {state.query}"

    try:
        result = call_with_fallback(prompt, task_type="general")
        answer = result.get("response", "")
        model = result.get("tier", "none")
    except Exception as exc:
        logger.error("rag_agent all tiers failed: %s", exc)
        answer = f"I found relevant context but couldn't generate a response: {exc}"
        model = "none"

    _telemetry_hook(state, answer, model, start)
    return {
        "kb_context": kb_context,
        "response": answer,
        "next_node": "formatter",
    }


def code_agent(state: MetisState) -> dict:
    """
    Generate code using fallback chain (Ollama → Opencode → Qwen → Gemini → Telegram).
    Validates any Python code blocks with ast.parse before returning.
    """
    start = time.monotonic()
    try:
        result = call_with_fallback(state.query, task_type="code")
        answer = result.get("response", "")
        model = result.get("tier", "none")
    except Exception as exc:
        logger.error("code_agent all tiers failed: %s", exc)
        answer = f"Code generation failed: {exc}"
        model = "none"

    validated_answer = _validate_code_blocks(answer)
    _telemetry_hook(state, validated_answer, model, start)

    return {
        "response": validated_answer,
        "code_snippet": _extract_code_block(validated_answer),
        "next_node": "formatter",
    }


def search_agent(state: MetisState, resume: bool = False) -> dict:
    """
    Web search via DuckDuckGo. Depth=1 (5 queries) for fresh searches.
    If resume=True, continues from stored search_context with 2 more layers.

    Returns answer + source citations + "go deeper" prompt.
    """
    start = time.monotonic()

    if resume:
        # Resume from previous search context
        try:
            prev = json.loads(state.search_context) if state.search_context else {}
        except json.JSONDecodeError:
            prev = {}
        return _search_recursive(state, prev, depth=2, start=start)
    else:
        # Fresh search: depth=1
        return _search_recursive(state, {}, depth=1, start=start)


def general_agent(state: MetisState) -> dict:
    """General-purpose fallback agent."""
    start = time.monotonic()
    result = _general_answer(state, start)
    _telemetry_hook(state, result.get("response", ""), settings.GENERAL_MODEL, start)
    return result


def file_reader_agent(state: MetisState) -> dict:
    """
    Safely read a file requested by the user.
    Validates path, checks permissions, extracts content (text, PDF, DOCX),
    formats for response with 3500-char truncation for Telegram.
    """
    start = time.monotonic()

    # Get path from state (set by router) or extract from query
    raw_path = state.file_path or _extract_path_from_query(state.query)

    if not raw_path:
        response = "📭 No pude detectar una ruta de archivo en tu mensaje. Intenta escribir la ruta completa, por ejemplo:\n\n`Lee el archivo config.py en ~/Vscode-projects/Metis`"
        _telemetry_hook(state, response, "file_reader", start, error="No path detected")
        return {"response": response, "next_node": "formatter"}

    # Validate
    resolved, error = _validate_file_path(raw_path)
    if error:
        _telemetry_hook(state, error, "file_reader", start, error=error)
        return {"response": error, "next_node": "formatter"}

    # Extract content based on file type
    ext = resolved.suffix.lower()

    # Intelligence: detect large PDFs and suggest RAG ingestion
    file_size = resolved.stat().st_size
    if ext == ".pdf" and file_size > 1_048_576:  # >1MB
        size_mb = file_size / 1_048_576
        response = (
            f"📄 **{resolved.name}** (`{resolved}`)\n"
            f"  Size: {size_mb:.1f}MB\n\n"
            "Este archivo es grande. ¿Prefieres que lo indexe en el RAG con "
            "`python -m src.memory.ingest` para que me hagas preguntas específicas? "
            "Así no tienes que esperar a que lea todo el contenido."
        )
        _telemetry_hook(state, response, "file_reader", start)
        return {"response": response, "file_path": str(resolved), "next_node": "formatter"}

    try:
        content = _extract_file_content(resolved, ext)
    except UnicodeDecodeError:
        response = f"⛔ No puedo leer `{resolved.name}` — parece ser un archivo binario sin soporte de extracción."
        _telemetry_hook(state, response, "file_reader", start, error="Binary file (no extractor)")
        return {"response": response, "file_path": str(resolved), "next_node": "formatter"}
    except PermissionError:
        response = f"⛔ Permission denied: `{resolved}`"
        _telemetry_hook(state, response, "file_reader", start, error="Permission denied")
        return {"response": response, "file_path": str(resolved), "next_node": "formatter"}
    except Exception as exc:
        logger.exception("file_reader_agent extraction failed for '%s': %s", resolved, exc)
        response = f"⚠️ Error leyendo archivo: {exc}"
        _telemetry_hook(state, response, "file_reader", start, error=str(exc))
        return {"response": response, "file_path": str(resolved), "next_node": "formatter"}

    # Format response
    size_kb = len(content.encode("utf-8")) / 1024
    lines = content.count("\n") + 1
    ext_label = ext.lstrip(".") or "text"

    header = f"📄 **{resolved.name}** (`{resolved}`)\n"
    header += f"  Size: {size_kb:.1f}KB | Lines: {lines} | Type: {ext_label}\n\n"

    # Truncate to 3500 chars max for Telegram limit
    if len(content) > 3500:
        preview = content[:3500]
        footer = f"\n\n...(truncated, showing first 3500 chars of {len(content)} total)"
        response = header + f"```\n{preview}\n```" + footer
    else:
        response = header + f"```\n{content}\n```"

    _telemetry_hook(state, response, "file_reader", start)
    return {
        "response": response,
        "file_path": str(resolved),
        "file_content": content,
        "next_node": "formatter",
    }


def _extract_file_content(path: Path, ext: str) -> str:
    """
    Extract text content from a file based on its extension.
    Supports: plain text, .pdf (PyMuPDF fitz), .docx (python-docx).
    Raises UnicodeDecodeError for unsupported binary files.
    """
    if ext == ".pdf":
        return _extract_pdf(path)
    elif ext == ".docx":
        return _extract_docx(path)
    else:
        # Plain text — let read_text raise UnicodeDecodeError for binaries
        return path.read_text(encoding="utf-8")


def _extract_pdf(path: Path) -> str:
    """
    Extract text from PDF using PyMuPDF (fitz).
    Falls back to pypdf if fitz is unavailable.
    """
    # Try PyMuPDF first (preferred — faster, better extraction)
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(str(path), filetype="pdf")
        pages = []
        for page_num in range(len(doc)):
            page = doc[page_num]
            text = page.get_text()
            pages.append(text)
        doc.close()
        full_text = "\n\n".join(pages)
        if full_text.strip():
            return full_text
        # If fitz extracted empty text, fall through to pypdf
    except ImportError:
        logger.debug("PyMuPDF (fitz) not available, trying pypdf for '%s'", path)
    except Exception as exc:
        logger.warning("PyMuPDF extraction failed for '%s': %s, falling back to pypdf", path, exc)

    # Fallback: pypdf
    try:
        from pypdf import PdfReader
        reader = PdfReader(str(path))
        pages = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages.append(text)
        full_text = "\n\n".join(pages)
        if full_text.strip():
            return full_text
    except ImportError:
        pass
    except Exception as exc:
        logger.error("pypdf also failed for '%s': %s", path, exc)

    raise ValueError(f"No se pudo extraer texto del PDF `{path.name}`. Puede estar escaneado o protegido.")


def _extract_docx(path: Path) -> str:
    """
    Extract text from .docx using python-docx.
    Handles paragraphs and tables.
    """
    try:
        from docx import Document
        doc = Document(str(path))
        parts = []

        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)

        # Extract tables (often contain important data)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text for cell in row.cells)
                if row_text.strip():
                    parts.append(row_text)

        full_text = "\n".join(parts)
        if full_text.strip():
            return full_text
    except ImportError:
        raise ImportError("python-docx no está instalado. Instala con: pip install python-docx")
    except Exception as exc:
        raise ValueError(f"Error extrayendo texto de `{path.name}`: {exc}")

    raise ValueError(f"No se pudo extraer texto del archivo DOCX `{path.name}`.")


def file_editor_agent(state: MetisState) -> dict:
    """
    Edit a file. Supports: find/replace, append, insert at line.
    First call asks for confirmation, second call (after user says yes) executes.
    """
    start = time.monotonic()

    raw_path = state.file_path or _extract_path_from_query(state.query)
    if not raw_path:
        response = "📭 No pude detectar qué archivo editar. Especifica la ruta, por ejemplo:\n\n`Edita config.py en ~/Vscode-projects/Metis y cambia 'debug=False' por 'debug=True'`"
        _telemetry_hook(state, response, "file_editor", start, error="No path detected")
        return {"response": response, "next_node": "formatter"}

    resolved, error = _validate_file_path(raw_path)
    if error:
        _telemetry_hook(state, error, "file_editor", start, error=error)
        return {"response": error, "next_node": "formatter"}

    # If awaiting confirmation, execute the edit
    if state.awaiting_confirmation and state.pending_action:
        import json
        try:
            action = json.loads(state.pending_action)
        except json.JSONDecodeError:
            response = "⚠️ Error interno: acción pendiente corrupta."
            _telemetry_hook(state, response, "file_editor", start, error="Corrupt pending action")
            return {"response": response, "awaiting_confirmation": False, "next_node": "formatter"}

        try:
            content = resolved.read_text(encoding="utf-8")
            original = content

            if action["type"] == "replace":
                if action["old"] not in content:
                    response = f"🔍 No encontré `{action['old'][:80]}...` en el archivo para reemplazar."
                    _telemetry_hook(state, response, "file_editor", start, error="Text not found")
                    return {"response": response, "file_path": str(resolved), "awaiting_confirmation": False, "next_node": "formatter"}
                content = content.replace(action["old"], action["new"], 1)
            elif action["type"] == "append":
                content += action["text"]
            elif action["type"] == "insert":
                lines = content.split("\n")
                line_idx = min(action["line"], len(lines))
                lines.insert(line_idx, action["text"])
                content = "\n".join(lines)

            resolved.write_text(content, encoding="utf-8")
            response = f"✅ **Archivo editado:** `{resolved.name}`\n\n"
            response += f"Ruta: `{resolved}`\n\n"
            response += f"**Cambios aplicados:**\n"
            if action["type"] == "replace":
                response += f"  Reemplazado: `{action['old'][:50]}...` → `{action['new'][:50]}...`\n"
            elif action["type"] == "append":
                response += f"  Agregado {len(action['text'])} chars al final\n"
            elif action["type"] == "insert":
                response += f"  Insertado en línea {action['line']}\n"
            response += f"\nTamaño final: {len(content.encode('utf-8')) / 1024:.1f}KB"

            _telemetry_hook(state, response, "file_editor", start)
            return {
                "response": response,
                "file_path": str(resolved),
                "file_content": content,
                "awaiting_confirmation": False,
                "pending_action": "",
                "next_node": "formatter",
            }
        except Exception as exc:
            # Restore original on error
            try:
                resolved.write_text(original, encoding="utf-8")
            except Exception:
                pass
            response = f"⚠️ Error editando archivo: {exc}"
            _telemetry_hook(state, response, "file_editor", start, error=str(exc))
            return {"response": response, "file_path": str(resolved), "awaiting_confirmation": False, "next_node": "formatter"}

    # First call: extract edit intent and ask for confirmation
    query_lower = state.query.lower()

    # Try to detect find/replace pattern
    import re
    replace_match = re.search(r"(?:cambia|replace|change|reemplaza|sustituye).+?['\"](.+?)['\"].+?['\"](.+?)['\"]", query_lower)
    if replace_match:
        old_text = replace_match.group(1)
        new_text = replace_match.group(2)
    else:
        # Try simpler: look for "from X to Y" or "por" pattern
        replace_match = re.search(r"['\"](.+?)['\"].+(?:por|to|→|->).+?['\"](.+?)['\"]", query_lower)
        if replace_match:
            old_text = replace_match.group(1)
            new_text = replace_match.group(2)
        else:
            old_text = ""
            new_text = ""

    if old_text and new_text:
        import json
        pending = json.dumps({"type": "replace", "old": old_text, "new": new_text})
        response = (
            f"✏️ **Editar archivo:** `{resolved.name}`\n\n"
            f"📋 Acción: Reemplazar\n"
            f"  De: `{old_text[:100]}`\n"
            f"  A: `{new_text[:100]}`\n\n"
            f"¿Confirmas este cambio? Responde **sí** o **no**."
        )
        _telemetry_hook(state, response, "file_editor", start)
        return {
            "response": response,
            "file_path": str(resolved),
            "awaiting_confirmation": True,
            "pending_action": pending,
            "route": ROUTE_FILE_EDIT,
            "next_node": ROUTE_FILE_EDIT,
        }

    # Generic edit: ask what to change
    response = (
        f"✏️ **Editar archivo:** `{resolved.name}` (`{resolved}`)\n\n"
        f"¿Qué cambio quieres hacer? Puedo:\n\n"
        f"1️⃣ **Reemplazar** texto: `cambia 'X' por 'Y'`\n"
        f"2️⃣ **Agregar** al final: `agrega esto al final: ...`\n"
        f"3️⃣ **Insertar** en línea: `inserta en línea 5: ...`\n\n"
        f"Describe el cambio que necesitas."
    )
    _telemetry_hook(state, response, "file_editor", start)
    return {"response": response, "file_path": str(resolved), "next_node": "formatter"}


def file_deleter_agent(state: MetisState) -> dict:
    """
    Delete a file. ALWAYS asks for confirmation first (2-step flow).
    """
    start = time.monotonic()

    raw_path = state.file_path or _extract_path_from_query(state.query)
    if not raw_path:
        response = "📭 No pude detectar qué archivo borrar. Especifica la ruta completa."
        _telemetry_hook(state, response, "file_deleter", start, error="No path detected")
        return {"response": response, "next_node": "formatter"}

    resolved, error = _validate_file_path(raw_path)
    if error:
        _telemetry_hook(state, error, "file_deleter", start, error=error)
        return {"response": error, "next_node": "formatter"}

    # If awaiting confirmation, execute the delete
    if state.awaiting_confirmation:
        try:
            resolved.unlink()
            response = f"🗑️ **Archivo eliminado:** `{resolved.name}`\n\nRuta: `{resolved.parent}`"
            _telemetry_hook(state, response, "file_deleter", start)
            return {
                "response": response,
                "file_path": str(resolved),
                "awaiting_confirmation": False,
                "pending_action": "",
                "next_node": "formatter",
            }
        except Exception as exc:
            response = f"⚠️ Error eliminando archivo: {exc}"
            _telemetry_hook(state, response, "file_deleter", start, error=str(exc))
            return {"response": response, "file_path": str(resolved), "awaiting_confirmation": False, "next_node": "formatter"}

    # First call: ask for confirmation
    import json
    pending = json.dumps({"type": "delete", "path": str(resolved)})
    size_kb = resolved.stat().st_size / 1024

    response = (
        f"⚠️ **Eliminar archivo**\n\n"
        f"📄 Archivo: `{resolved.name}`\n"
        f"📍 Ruta: `{resolved}`\n"
        f"📏 Tamaño: {size_kb:.1f}KB\n\n"
        f"¿Estás seguro de que quieres **eliminar permanentemente** este archivo?\n"
        f"Responde **sí** para confirmar o **no** para cancelar."
    )
    _telemetry_hook(state, response, "file_deleter", start)
    return {
        "response": response,
        "file_path": str(resolved),
        "awaiting_confirmation": True,
        "pending_action": pending,
        "route": ROUTE_FILE_DELETE,
        "next_node": ROUTE_FILE_DELETE,
    }


def bash_agent(state: MetisState) -> dict:
    """
    Execute a safe bash command. Validates against whitelist and dangerous patterns.
    """
    import subprocess

    start = time.monotonic()
    query = state.query.strip()

    # Extract the command: look for patterns like "run X", "execute X", or just the command itself
    query_lower = query.lower()
    cmd = query

    # Strip command prefixes
    for prefix in ["run", "run the command", "execute", "ejecuta", "corre", "correr", "bash", "shell"]:
        if query_lower.startswith(prefix):
            cmd = query[len(prefix):].strip().strip('"').strip("'").strip()
            break

    if not cmd:
        response = "⚠️ No detecté un comando para ejecutar. Ejemplo:\n\n`Ejecuta ls -la ~/Vscode-projects/Metis/src`"
        _telemetry_hook(state, response, "bash", start, error="No command detected")
        return {"response": response, "next_node": "formatter"}

    # Validate: first word must be in whitelist
    parts = cmd.split()
    base_cmd = parts[0] if parts else ""

    if base_cmd not in settings.BASH_ALLOWED_COMMANDS:
        response = f"⛔ Comando `{base_cmd}` no permitido. Comandos disponibles:\n\n`ls`, `cat`, `head`, `tail`, `wc`, `find`, `grep`, `du`, `df`, `pwd`, `tree`, `stat`, `diff`, `sort`, `uniq`, `echo`, `date`, `whoami`, `uname`, `uptime`"
        _telemetry_hook(state, response, "bash", start, error=f"Command not allowed: {base_cmd}")
        return {"response": response, "next_node": "formatter"}

    # Check for dangerous patterns
    for pattern in settings.BASH_DANGEROUS:
        if pattern in cmd:
            response = f"⛔ Comando bloqueado por seguridad: contiene `{pattern}`"
            _telemetry_hook(state, response, "bash", start, error=f"Dangerous pattern: {pattern}")
            return {"response": response, "next_node": "formatter"}

    # Ensure working directory is under allowed root
    allowed_root = settings.ALLOWED_FILE_ROOT.resolve()

    # Execute
    try:
        proc = subprocess.run(
            cmd,
            shell=True,
            capture_output=True,
            text=True,
            timeout=settings.BASH_TIMEOUT,
            cwd=str(allowed_root),
        )

        stdout = (proc.stdout or "").strip()
        stderr = (proc.stderr or "").strip()

        # Truncate output if too large
        output = stdout if stdout else stderr
        if len(output) > settings.MAX_BASH_OUTPUT:
            output = output[:settings.MAX_BASH_OUTPUT] + f"\n\n...(truncated, showing first {settings.MAX_BASH_OUTPUT} chars)"

        if proc.returncode == 0:
            response = f"🖥️ **Command:** `{cmd}`\n\n```\n{output}\n```"
            _telemetry_hook(state, response, "bash", start)
            return {
                "response": response,
                "bash_output": output,
                "bash_error": "",
                "next_node": "formatter",
            }
        else:
            response = f"🖥️ **Command:** `{cmd}`\n\n⚠️ Exit code: {proc.returncode}\n\n```\n{output}\n```"
            _telemetry_hook(state, response, "bash", start, error=f"Exit {proc.returncode}")
            return {
                "response": response,
                "bash_output": output,
                "bash_error": stderr,
                "next_node": "formatter",
            }
    except subprocess.TimeoutExpired:
        response = f"⏱️ Timeout: el comando `{cmd}` tardó más de {settings.BASH_TIMEOUT}s"
        _telemetry_hook(state, response, "bash", start, error="Timeout")
        return {"response": response, "next_node": "formatter"}
    except Exception as exc:
        response = f"⚠️ Error ejecutando comando: {exc}"
        _telemetry_hook(state, response, "bash", start, error=str(exc))
        return {"response": response, "next_node": "formatter"}


def echo_agent(state: MetisState) -> dict:
    """
    Echo pronunciation practice agent.
    
    Flow:
    1. Get target sentence from database
    2. Send to user with instruction to read it
    3. (Next message) Transcribe user's voice
    4. Score pronunciation
    5. Generate TTS of correct version
    6. Return feedback
    """
    from src.echo import EchoScorer, WhisperSTT, EchoTTS, EchoDatabase
    
    start = time.monotonic()
    
    try:
        db = EchoDatabase()
        scorer = EchoScorer()
        stt = WhisperSTT()
        tts = EchoTTS()
        
        # Check if user specified a custom sentence
        query = state.query.lower().strip()
        
        # Extract custom sentence from "practice: <text>" or "repeat: <text>"
        custom_sentence = None
        for prefix in ["practice:", "repeat:", "practicar:", "repetir:", "say this:", "lee esto:"]:
            if prefix in query:
                custom_sentence = query.split(prefix, 1)[1].strip()
                break
        
        # Get sentence from database
        if custom_sentence:
            target = custom_sentence
            language = "en"  # Default, could be detected
        else:
            # Extract level from query (default: A1)
            level = "A1"
            for lvl in ["A1", "A2", "B1", "B2", "C1"]:
                if lvl.lower() in query:
                    level = lvl
                    break
            
            # Extract language from query
            language = "en"
            if "spanish" in query or "español" in query or "en español" in query:
                language = "es"
            
            sentence_data = db.get_sentence(level=level, language=language)
            
            if not sentence_data:
                response = f"📭 No sentences found for level {level} in {language.upper()}. Try a different level or language."
                _telemetry_hook(state, response, "echo", start, error="No sentences found")
                return {"response": response, "next_node": "formatter"}
            
            target = sentence_data["text"]
            language = sentence_data["language"]
        
        # Send target sentence to user
        response = f"📖 **Read this:**\n\n*{target}*\n\nNow send a voice message reading it aloud! 🎤"
        
        # For now, we'll return the target sentence and let the Telegram handler
        # manage the voice message flow. In a full implementation, this would
        # set state.awaiting_echo_voice = True and wait for the next message.
        
        _telemetry_hook(state, response, "echo", start)
        return {
            "response": response,
            "next_node": "formatter",
        }
        
    except Exception as exc:
        logger.exception(f"Echo agent failed: {exc}")
        response = f"⚠️ Echo practice failed: {exc}"
        _telemetry_hook(state, response, "echo", start, error=str(exc))
        return {"response": response, "next_node": "formatter"}


def formatter(state: MetisState) -> dict:
    """Final formatting node. Ensures response is clean and presentable."""
    response = state.response.strip()
    if not response:
        response = "I couldn't generate a response for that query. Please try rephrasing."

    # Append "go deeper" hint for search responses
    if state.route == ROUTE_SEARCH and state.search_context:
        response += "\n\n💡 Reply with *deeper* for more detailed research."

    return {
        "response": response,
        "next_node": "__end__",
    }


# ---- internal helpers ----

def _general_answer(state: MetisState, start: float, chat_id: int | None = None) -> dict:
    try:
        result = call_with_fallback(state.query, task_type="general", chat_id=chat_id)
        answer = result.get("response", "")
        model = result.get("tier", "fallback")
    except Exception as exc:
        logger.error("general_agent all tiers failed: %s", exc)
        answer = f"Sorry, I couldn't process your request: {exc}"
        model = "none"
    _telemetry_hook(state, answer, model, start)
    return {
        "response": answer,
        "next_node": "formatter",
    }


def _validate_code_blocks(text: str) -> str:
    """Try ast.parse on every Python code block found. Append validation notes."""
    import re
    blocks = re.findall(r"```python\n(.*?)\n```", text, re.DOTALL)
    notes: list[str] = []
    for i, block in enumerate(blocks):
        try:
            ast.parse(block)
        except SyntaxError as se:
            notes.append(f"⚠️ Code block {i+1} has a syntax error: {se}")
    if notes:
        text += "\n\n" + "\n".join(notes)
    return text


def _extract_code_block(text: str) -> str:
    import re
    match = re.search(r"```(?:python)?\n(.*?)\n```", text, re.DOTALL)
    return match.group(1) if match else ""


# ---- search engine ----

def _duckduckgo_search(query: str, max_results: int = 5) -> list[dict[str, str]]:
    """Run a DuckDuckGo search, return list of {title, snippet, url}."""
    try:
        from duckduckgo_search import DDGS
        with DDGS() as ddgs:
            results = list(ddgs.text(query, max_results=max_results))
        return [
            {"title": r.get("title", ""), "snippet": r.get("body", ""), "url": r.get("href", "")}
            for r in results
        ]
    except Exception as exc:
        logger.error("DuckDuckGo search failed for '%s': %s", query[:80], exc)
        return []


def _search_recursive(state: MetisState, prev_context: dict, depth: int, start: float) -> dict:
    """
    Recursive search loop.

    1. LLM generates search queries from current context
    2. DuckDuckGo searches each query
    3. LLM synthesizes learnings
    4. If depth > 0, recurse
    """
    model = settings.GENERAL_MODEL
    learnings = prev_context.get("learnings", [])
    all_sources = prev_context.get("sources", [])
    query_history = prev_context.get("queries", [])

    # Step 1: Generate search queries
    if not query_history or depth > 1:
        # Fresh queries generation
        system_prompt = (
            "You are a research assistant. Generate 3-5 specific search queries "
            "to thoroughly research the topic. Return ONLY a JSON array of strings."
        )
        context_note = f"\nPrevious findings: {learnings}" if learnings else ""
        prompt = f"Topic: {state.query}{context_note}\nGenerate search queries:"

        try:
            result = call_with_fallback(prompt, task_type="general")
            queries_raw = result.get("response", "")
            # Parse JSON from response
            import re
            json_match = re.search(r'\[.*?\]', queries_raw, re.DOTALL)
            if json_match:
                queries = json.loads(json_match.group(0))
            else:
                queries = [state.query]
        except Exception:
            queries = [state.query]

        if not isinstance(queries, list):
            queries = [str(queries)]
    else:
        queries = query_history[:3]

    # Step 2: Search each query
    new_results: list[dict[str, str]] = []
    for q in queries[:5]:  # cap at 5
        results = _duckduckgo_search(q, max_results=3)
        new_results.extend(results)

    if not new_results and not learnings:
        # Total failure
        _telemetry_hook(state, "No search results found.", model, start)
        return {
            "response": f"I searched the web for '{state.query}' but found no relevant results.",
            "search_context": json.dumps({"learnings": [], "sources": [], "queries": queries}),
            "next_node": "formatter",
        }

    # Step 3: Synthesize learnings from results
    snippets = "\n".join(
        f"- [{r['title']}]({r['url']}): {r['snippet']}"
        for r in new_results
    )
    all_sources.extend(new_results)

    system_prompt = (
        "You are a research analyst. Read the search results below and extract "
        "key findings as a bullet list. Be factual and concise."
    )
    prev_findings = f"\nPrevious findings: {learnings}" if learnings else ""
    prompt = f"{prev_findings}\n\nNew search results:\n{snippets}\n\nSynthesize key findings:"

    try:
        result = call_with_fallback(prompt, task_type="general")
        new_learnings = result.get("response", "")
    except Exception as exc:
        new_learnings = f"Error synthesizing: {exc}"

    all_learnings = learnings + [new_learnings] if learnings else [new_learnings]

    # Step 4: Recurse or finalize
    if depth > 1:
        # Continue the loop
        return _search_recursive(
            state,
            {"learnings": all_learnings, "sources": all_sources, "queries": queries},
            depth=depth - 1,
            start=start,
        )

    # Final: compile report
    sources_text = "\n".join(
        f"  - [{r['title']}]({r['url']})" for r in all_sources[:10]
    )
    findings = "\n\n".join(all_learnings)

    report = f"## Research: {state.query}\n\n{findings}"
    if sources_text:
        report += f"\n\n### Sources\n{sources_text}"

    report += "\n\n💡 Reply with *deeper* for more detailed research."

    _telemetry_hook(state, report, model, start)

    return {
        "response": report,
        "kb_context": findings,
        "search_context": json.dumps({
            "learnings": all_learnings,
            "sources": all_sources,
            "queries": queries,
        }),
        "next_node": "formatter",
    }


# --- Research agent (vault builder) ---

ROUTE_RESEARCH = "research"

_RESEARCH_KEYWORDS = {
    "research", "investiga", "investigar", "study", "estudio",
    "build vault", "crear nota", "create note", "obsidian",
    "knowledge base", "base de conocimiento", "wiki",
}


def research_agent(state: MetisState) -> dict:
    """
    Research a topic and build knowledge in the Obsidian vault.

    Karpathy pattern:
    1. Search the web for the topic (DuckDuckGo)
    2. Save raw content to vault/raw/
    3. LLM synthesizes structured notes from raw content
    4. Writes notes to vault/wiki/ with frontmatter, wikilinks, indexes
    5. Reports what was created

    Returns dict compatible with StateGraph merge.
    """
    start = time.monotonic()
    query = state.query.strip()

    # Step 1: Web search to gather raw content
    raw_sources = _duckduckgo_search(query, max_results=8)
    if not raw_sources:
        # No web results — still create a note from the query itself
        raw_content = f"User query: {query}\n\nNo web sources found. Creating note from domain knowledge."
    else:
        raw_content = "\n\n".join(
            f"## {r['title']}\nURL: {r['url']}\n\n{r['snippet']}"
            for r in raw_sources[:6]
        )

    # Step 2: Save raw content to vault
    from src.vault import save_raw
    try:
        raw_path = save_raw(query, raw_content)
    except Exception as exc:
        logger.warning("Failed to save raw content: %s", exc)
        raw_path = None

    # Step 3: LLM classifies and structures notes
    classify_prompt = (
        f"Analyze this research topic and determine how to structure the knowledge.\n\n"
        f"Topic: {query}\n\n"
        f"Raw content:\n{raw_content[:3000]}\n\n"
        f"Return a JSON array of 2-4 notes to create. Each note must have:\n"
        f"- title: string (clear, concise, Title Case)\n"
        f"- type: one of [concept, tool, pattern, project]\n"
        f"- domain: one of [ai, devops, security, web, math]\n"
        f"- tags: array of 3-6 tags (no # prefix)\n"
        f"- body: markdown body with ## sections (at least 200 words)\n"
        f"- summary: one-line summary for index tables\n"
        f"- related: array of 1-3 existing note titles this connects to (or empty)\n\n"
        f"Return ONLY a valid JSON array. No markdown code blocks, no explanations."
    )

    try:
        result = call_with_fallback(classify_prompt, task_type="general")
        notes_raw = result.get("response", "[]")
        # Parse JSON from response
        json_match = re.search(r'\[.*\]', notes_raw, re.DOTALL)
        if json_match:
            notes = json.loads(json_match.group(0))
        else:
            notes = []
    except Exception as exc:
        logger.error("research_agent classification failed: %s", exc)
        notes = []

    if not notes:
        # Fallback: create a single general note
        notes = [{
            "title": query[:80].title(),
            "type": "note",
            "domain": "ai",
            "tags": ["research"],
            "body": f"## Overview\n\nResearch on: {query}\n\nNo structured notes could be generated from available sources.",
            "summary": f"Research: {query[:100]}",
            "related": [],
        }]

    # Step 4: Write notes to vault
    from src.vault import write_notes_batch
    try:
        created_paths = write_notes_batch(notes)
    except Exception as exc:
        logger.error("research_agent vault write failed: %s", exc)
        created_paths = []

    # Step 5: Build response
    if created_paths:
        vault_root = os.getenv("METIS_VAULT_ROOT", "~/.metis/vault")
        note_list = "\n".join(
            f"  📝 **{n.get('title', 'Unknown')}** → `{p}`"
            for n, p in zip(notes, created_paths)
        )
        response = (
            f"📚 **Research vault updated:** {len(created_paths)} note(s) created\n\n"
            f"{note_list}\n\n"
            f"📂 Vault: `{vault_root}`\n"
            f"💡 Open in Obsidian to see wikilinks and graph view."
        )
    else:
        response = "⚠️ No notes could be generated from available sources."

    _telemetry_hook(state, response, "research_vault", start)
    return {
        "response": response,
        "next_node": "formatter",
    }


# --- Routing function for LangGraph conditional edges ---

def route_decision(state: MetisState) -> str:
    """Used as the callable in `graph.add_conditional_edges('router', route_decision)`."""
    if not state.route:
        result = router(state)
        return result.get("next_node", ROUTE_GENERAL)
    return state.route
